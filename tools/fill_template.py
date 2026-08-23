from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

doc = Document(r"D:\1.个人资料\彼岸墓园算法开发上线资料\落实算法安全主体责任基本情况模板.docx")

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text=""):
    return doc.add_paragraph(text)

def add_bold_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_table_row(table, key, value):
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = value

# 清空现有内容（保留第一个段落作为参考）
# 实际上我们保留原格式，在后面添加内容

# 在文档末尾添加内容
content_sections = [
    ("一、算法安全专职机构", 1),
    ("（一）算法安全专职机构设置", 2),
]

# 找到最后一个段落的位置
paras = list(doc.paragraphs)

# 添加标题
doc.add_heading("【彼岸墓园 · 填好的版本】", 0)
doc.add_paragraph("以下为草稿内容，请刘贻清核对并替换【】内字段后填入模板。")
doc.add_paragraph("")

# 一、算法安全专职机构
doc.add_heading("一、算法安全专职机构", 1)

doc.add_heading("（一）算法安全专职机构设置", 2)
doc.add_paragraph("算法安全专职机构名称：算法安全管理小组（虚拟组织）")
doc.add_paragraph("")
doc.add_paragraph("组织架构：")
doc.add_paragraph("• 组长（算法安全责任人）：【刘贻清】——公司法定代表人，统筹算法安全工作")
doc.add_paragraph("• 组员：【暂无专职人员，由法人兼任日常管理职责】")
doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("注：")
run.italic = True
p.add_run("本公司为初创小微企业，暂无专职算法安全团队。算法安全主体责任由法定代表人直接承担。后续随业务发展，将适时设立专职算法安全岗位。")

doc.add_heading("（二）算法安全专职机构负责人基本信息", 2)
table = doc.add_table(rows=1, cols=2)
hdr = table.rows[0].cells
hdr[0].text = "项目"
hdr[1].text = "内容"
rows_data = [
    ("姓名", "【刘贻清】"),
    ("职务", "法定代表人、执行董事"),
    ("身份证号", "【待填】"),
    ("联系电话", "【待填】"),
    ("邮箱", "【待填】"),
    ("职责", "全面负责算法安全工作，承担算法安全主体责任"),
]
for k, v in rows_data:
    add_table_row(table, k, v)

doc.add_paragraph("")
doc.add_heading("（三）算法安全工作人员任职要求", 2)
doc.add_paragraph("由于公司现为1人在职，算法安全相关工作由法定代表人直接承担。未来扩充团队后，拟设立以下任职要求：")
doc.add_paragraph("• 具有计算机或信息安全相关背景")
doc.add_paragraph("• 熟悉《深度合成管理规定》《算法推荐管理规定》等法规")
doc.add_paragraph("• 具备内容安全审核经验")

doc.add_heading("（四）算法安全工作人员配备规模", 2)
doc.add_paragraph("当前：1人（法定代表人兼任）")
doc.add_paragraph("计划：随用户规模扩大，逐步配置专职内容审核人员")

doc.add_heading("（五）算法安全技术保障措施", 2)
table2 = doc.add_table(rows=1, cols=2)
hdr2 = table2.rows[0].cells
hdr2[0].text = "措施"
hdr2[1].text = "具体方式"
tech_rows = [
    ("内容安全审核", "接入阿里云内容安全 API，对用户输入文本进行自动审核"),
    ("人工审核兜底", "AI 生成内容进入人工审核队列后方可公开展示"),
    ("日志留存", "操作日志留存不少于6个月"),
    ("数据备份", "每日自动备份数据库"),
]
for k, v in tech_rows:
    add_table_row(table2, k, v)

doc.add_paragraph("")

# 二、算法安全管理制度建设
doc.add_heading("二、算法安全管理制度建设", 1)

doc.add_heading("（一）算法安全自评估制度建设", 2)
doc.add_paragraph("本公司参照《互联网信息服务深度合成管理规定》，建立算法安全自评估制度，在以下时点开展自评估：")
doc.add_paragraph("1. 新算法能力上线前")
doc.add_paragraph("2. 重大功能变更时")
doc.add_paragraph("3. 定期评估（每年至少一次）")
doc.add_paragraph("")
doc.add_paragraph("现状说明：当前算法均调用第三方基础模型（火山方舟 API），本公司不进行模型训练与微调。自评估报告已编写并存档。")
doc.add_paragraph("执行保障措施：法定代表人负责自评估执行，评估结果记录存档，发现问题立即整改。")

doc.add_heading("（二）算法安全监测制度建设", 2)

doc.add_paragraph("1. 信息安全监测")
doc.add_paragraph("制度：用户输入文本（祭品描述、留言、时间线等）均经内容安全 API 自动审核，违规内容自动拦截或进入人工审核队列。")
doc.add_paragraph("技术保障：阿里云内容安全 API（comment_detection_pro 规则集），审核异常时进入人工审核，不直接展示。")

doc.add_paragraph("2. 数据安全监测")
doc.add_paragraph("制度：数据库存储于本地磁盘，传输全程 HTTPS 加密；用户可申请数据导出与删除。")
doc.add_paragraph("技术保障：better-sqlite3 数据库，WAL 模式；文件存储于隔离目录；删除纪念馆时级联清理全部关联素材。")

doc.add_paragraph("3. 用户个人信息安全监测")
doc.add_paragraph("制度：最小化收集原则；数字人生成强制近亲属授权声明，留存授权记录；用户可申请注销账号。")
doc.add_paragraph("技术保障：/legal/privacy 已发布隐私政策；data_requests 流程支持用户数据导出与删除。")

doc.add_paragraph("4. 算法安全监测")
doc.add_paragraph("制度：监控算法生成服务的可用性（API 超时/失败时自动降级）；监控异常调用行为（配额限制防止滥用）。")
doc.add_paragraph("技术保障：每用户每月免费生成配额 3 次（ai_quotas 表）；任务状态跟踪（pending → processing → reviewing → done/failed）。")

doc.add_heading("（三）算法安全事件应急处理制度建设", 2)
table3 = doc.add_table(rows=1, cols=3)
hdr3 = table3.rows[0].cells
hdr3[0].text = "场景"
hdr3[1].text = "处置步骤"
hdr3[2].text = "责任人"
emergency_rows = [
    ("AI 生成服务不可用（API 超时/故障）", "1. 自动标记任务失败\n2. 退还用户已支付额度\n3. 服务降级关闭生成入口", "【刘贻清】"),
    ("发现违法有害生成内容", "1. 管理员下架内容\n2. 溯源日志留存\n3. 依法向主管部门报告", "【刘贻清】"),
    ("用户数据泄露风险", "1. 立即停止相关服务\n2. 排查泄露原因\n3. 通知受影响用户\n4. 向主管部门报告", "【刘贻清】"),
]
for k, v, r in emergency_rows:
    row = table3.add_row()
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[2].text = r

doc.add_paragraph("")
p = doc.add_paragraph()
p.add_run("应急联络：").bold = True
p.add_run("算法安全责任人【刘贻清】，电话【待填】，24小时可达")

doc.add_heading("（四）算法违法违规处置制度建设", 2)
table4 = doc.add_table(rows=1, cols=3)
hdr4 = table4.rows[0].cells
hdr4[0].text = "违规类型"
hdr4[1].text = "具体情形"
hdr4[2].text = "处置方式"
violation_rows = [
    ("数据使用违规", "超出隐私政策范围使用用户数据", "立即停止违规行为，删除相关数据"),
    ("信息安全违规", "用户上传违规内容（经审核发现）", "拒绝展示，情节严重报告主管部门"),
    ("用户权益保护违规", "未履行知情权/删除权", "补充完善功能，满足用户诉求"),
    ("算法安全违规", "AI 生成内容出现安全问题", "下架内容，排查模型调用链路"),
]
for k, v, r in violation_rows:
    add_table_row(table4, k, f"{v} → {r}")

doc.add_paragraph("")
doc.add_heading("（五）其他制度", 2)
doc.add_paragraph("本公司目前暂无其他补充制度，后续将根据业务发展补充完善。")

# 三、附件
doc.add_heading("三、附件", 1)
doc.add_paragraph("1. 算法安全自评估报告（已有）")
doc.add_paragraph("2. 隐私政策（/legal/privacy，已上线）")
doc.add_paragraph("3. 用户协议（/legal/terms，已上线）")

# 刘贻清待填字段
doc.add_paragraph("")
doc.add_heading("刘贻清待填字段清单", 2)
doc.add_paragraph("• 公司全称：【待填】")
doc.add_paragraph("• 法定代表人/算法安全责任人姓名：【刘贻清】")
doc.add_paragraph("• 身份证号：【待填】")
doc.add_paragraph("• 联系电话：【待填】")
doc.add_paragraph("• 邮箱：【待填】")
doc.add_paragraph("• 注册地址：【待填】")

output_path = r"D:\1.个人资料\彼岸墓园算法开发上线资料\落实算法安全主体责任基本情况-草稿.docx"
doc.save(output_path)
print(f"已保存到: {output_path}")
